# 导入所需的库
import os
import glob
from docx import Document

# 读取一个已存在的word文档对象
doc = Document("中药推荐软件部分源码.docx") # 你可以修改这个文件名

# 定义要查找的目录和文件后缀列表
directory = "kgweb" # 你可以修改这个目录
suffixes = [".py", ".j2", ".html", ".sql", ".sh", ".css", ".js"] # 你可以修改这个后缀列表

# 递归地遍历目录下所有符合后缀列表中任意一个的文件
for suffix in suffixes:
    for filename in glob.iglob(directory + '/**/*' + suffix, recursive=True):
        if ".min" in filename:
            continue
        # 打开文件并读取内容
        with open(filename, encoding="utf-8") as f:
            content = f.read()
        # 在word文档的末尾添加一个段落，写入文件名和内容
        doc.add_paragraph(filename)
        doc.add_paragraph(content)
        doc.add_paragraph("")

# 保存word文档
doc.save("output.docx") # 你可以修改这个文件名
